"""
Simple app to upload an image via a web form 
and view the inference results on the image in the browser.
"""
from subprocess import STDOUT, check_call , call,run
from datetime import datetime
from fastapi.responses import FileResponse
# import argparse
import io
import math
import os
from PIL import Image
import requests
import numpy as np
from base64 import b64encode
# import pythoncom
import re
from docx.enum.table import WD_ALIGN_VERTICAL

import base64
import torch


import pandas as pd

import openpyxl
# from flask import Flask, render_template, request, redirect,jsonify,send_file

import os
import json
# from docx2pdf import convert
from docx import Document # for pdf format
from docx.shared import Pt # for pdf format
from docx.shared import Inches


def predict(data_raw):

    github='ultralytics/yolov5'
    torch.hub.list(github, trust_repo=True)
    model = torch.hub.load("ultralytics/yolov5", "custom", path = "./rings18.pt", force_reload=True)
    
    model.classes=[3 ,10,11 ,12, 17]


    #print("Here")
    data_raw = data_raw
    data_raw1={'cylinder2': None, 'cylinder3': None, 'cylinder4': None, 'cylinder5': None, 'cylinder6': None, 'cylinder7': None, 'cylinder8': None, 'cylinder9': None, 'cylinder10': None, 'cylinder11': None, 'cylinder12': None, 'cylinder13': None, 'cylinder14': None, 'cylinder15': None, 
    'VESSEL_OBJECT_ID': 0,   'EQUIPMENT_CODE': 'string', 'EQUIPMENT_ID': 0, 'JOB_PLAN_ID': 0, 'JOB_ID': 0, 'LOG_ID': 0, 'Vessel': 'string', 'Hull_No': None,
    'Vessel_Type': None,'Local_Start_Time': None,'Time_Zone1': None, 'Local_End_Time': '2023-07-13 11:52:46','Time_Zone2': None,'Form_No': None, 'IMO_No': 'string', 
    'Maker': None, 'Model': None,'License_Builder': None,'Serial_No': None, 'MCR': None, 'Speed_at_MCR': None, 'Bore': None,'Stroke': None, 
    'Maker_T': None, 'Model_T': None, 'Total_Running_Hour': 'string','Cylinder_Oil_type': 'string','Normal_service_load_in_percentage_of_MCR': 'string', 
    'Scrubber': 'string', 'Position': 'string','ME_Cyllinder_Oil_Consumption': 'string', 'Me_Running_Hour_Since_Last_Check': 'string', 'Inspected_by_Rank': 'string', 
    'Fuel_Sulphur_percentage': 'string'}
    data_raw.update(data_raw1)
    data_raw['cylinder1']=data_raw["image"]
    del data_raw['image']
    # print("***********************************")
    # data=dict(itertools.islice(data_raw.items(), 8,len(data_raw)))
    # data_id1 =dict(itertools.islice(data_raw.items(), 4,8))
    # data_id=dict(itertools.islice(data_raw.items(), 0,4))
    
    l_raw=list(data_raw.keys())
    # print(l_raw)
    l_id=['VESSEL_OBJECT_ID','EQUIPMENT_CODE', 'EQUIPMENT_ID', 'JOB_PLAN_ID','JOB_ID','LOG_ID']
    l_pdf_para=['Vessel','Local_End_Time','IMO_No','Total_Running_Hour','Cylinder_Oil_type','Normal_service_load_in_percentage_of_MCR','Scrubber','Position','ME_Cyllinder_Oil_Consumption','Me_Running_Hour_Since_Last_Check','Inspected_by_Rank','Fuel_Sulphur_percentage']
    l_data=[]
    for i in l_raw:
        if i not in l_id and i not in l_pdf_para:
            l_data.append(i)
    data_id={}
    for i in l_id:
        if i not in data_id.keys():
            data_id[i]=data_raw[i]
    data_pdf_para={}
    for i in l_pdf_para:
        if i not in data_pdf_para.keys():
            data_pdf_para[i]=data_raw[i]
    data={}
    for i in l_data:
        if i not in data.keys():
            data[i]=data_raw[i]
    data_keys=list(data.keys())
    # print("data:",data)
    # print("data_keys:",data_keys)
    # print("data_id:",data_id)   
    # print("data_pdf_para:",data_pdf_para)
    cyl_index=-1
    defect_df_all_cyl={}
    img_list=[]
    obj_lst=[]
    cyl_obj={}
    data = {k: v for k, v in data.items() if v is not None}
    # print("data:",data.keys())
    data_k=list(data.keys())
    cyl_num_temp=[re.split(r'(\d+)', s) for s in data_k]
    print("cyl_num_temp",cyl_num_temp)
    cyl_num=[i[1] for i in cyl_num_temp]
    len_tensor=0
    cyl_pred={}

    for (cyl_number,cyl) in  zip(data.keys(),data.values()):
        cyl_index=cyl_index+1
#---
        header, encoded = cyl.split(",", 1)
        data = base64.b64decode(encoded)

        with open("image.png", "wb") as f:
            f.write(data)   

        img = Image.open(io.BytesIO(data))
        results = model(img, size=640)
        img = np.squeeze(results.render())
        # datatoexcel = pd.ExcelWriter('results.xlsx')
        # results.to_excel(datatoexcel)
        # datatoexcel.save()

        #print("RESULT=======",results)
        file_object = io.BytesIO()
        
        data = Image.fromarray(img)
        data.save(file_object, 'JPEG')
        
        base64img = "data:image/png;base64,"+b64encode(file_object.getvalue()).decode('ascii')

        res_tensor=results.xyxy[0]  # im1 predictions (tensor)
        # print("res_tensor:",res_tensor)
        len_tensor=len(res_tensor)
        if len(res_tensor)==0:
            img_list.append(file_object)
            cyl_pred[cyl_number]={"Overall_fault_id":"-99","Overall_Rating":"3","Recommendation":"(Satisfactory) - No fault - No recommendation"}
            defect_df_all_cyl["cylinder"+(data_keys[cyl_index])]={'lubrication': {'Ring1': '*', 'Ring2': '*', 'Ring3': '*', 'Ring4': '*'}, 
            'surface': {'Ring1': '*', 'Ring2': '*', 'Ring3': '*', 'Ring4': '*'}, 
            'deposits': {'Ring1': '*', 'Ring2': '*', 'Ring3': '*', 'Ring4': '*'}, 
            'breakage': {'Ring1': '*', 'Ring2': '*', 'Ring3': '*', 'Ring4': '*'}}
            obj_lst.append({'Fault_id':"999",'Rating':"3","Recommendation":"(Satisfactory) - No fault - No recommendation"})
                 
        else:
            h=data.height

        
    ################################################################################################################## Object #############################################################################################
            ##{"3"-C: '4',"12"-S: '3',"17"-LC: '5',"10"-O: '2', "11"-OB:'1'}
            faultdict = {"3": '4',"12": '3',"17": '5',"10": '2', "11":'1'}
            conf_lvl = {"0":"Unacceptable","1":"Marginal","2":"Fair","3":"Satisfactory"}
            
            #result coverted to Dataframe
            df = results.pandas().xyxy[0]  
            # print("dataframe:",df)

            df["rings"] = list(map(lambda x: "ring1" if x/h <=0.25 else ("ring2" if x/h <=0.45 else ("ring3" if x/h <=0.75 else "ring4")),df['ymin'])) #Map rings to rings column
            df["rating"] = list(map(lambda x: "Satisfactory" if x<0.25 else ("Fair" if x <0.5 else ("Marginal" if x<0.75 else "Unacceptable") ),df['confidence'])) #Map confidence level to conf column
            df["confidence_lvl"] = list(map(lambda x: "3" if x<0.25 else ("2" if x <0.5 else ("1" if x<0.75 else "0") ),df['confidence']))
            df["class"] = list(map(lambda x: "4" if x==3 else ("3" if x ==12 else ("5" if x==17 else ("2" if x==10 else "1")) ),df['class']))
            df["fault"] = list(map(lambda x: "Oil Black" if x=="1" else ("Too Much Oil" if x =="2" else ("Scratch" if x=="3" else ("Collapsed" if x=="4" else "Carbon")) ),df['class'])) #Map class to fault column
            # print("df:",df)
            #List of list of all parameters
            defect=[]
            for r in range(len(df)):
                ls=[]
                ls.append(df.loc[r]['rings'])
                ls.append(df.loc[r]["rating"])
                ls.append(df.loc[r]["class"])
                ls.append(df.loc[r]["fault"])
                defect.append(ls)
            print("defect_beforeloop:",defect)

            # print('len(defect):',len(defect))
            #Appending recommendations dummy values to defects list
            # for k in faultdict.values():
            #     for d in range(len(defect)):
            #         if defect[d][1]=="Unacceptable" and defect[d][2]==k:   
            #             defect[d].append("s"+str(k)+ str(list(conf_lvl.keys())[0]))
            #         elif defect[d][1]=="Marginal" and defect[d][2]==k:   
            #             defect[d].append("s"+str(k)+ str(list(conf_lvl.keys())[1]))
            #         elif defect[d][1]=="Fair" and defect[d][2]==k:   
            #             defect[d].append("s"+str(k)+ str(list(conf_lvl.keys())[2]))
            #         elif defect[d][1]=="Satisfactory" and defect[d][2]==k:   
            #             defect[d].append("s"+str(k)+ str(list(conf_lvl.keys())[3]))
            
            for k in faultdict.values():
                for d in range(len(defect)):
                    if defect[d][2]==k and k=="5":
                        defect[d].append("""1) Replace or Overhaul fuel injector to avoid improper combustion
                                    2) Adjust fuel temperature/viscocity to attain correct viscocity as per Maker's recommendation
                                    3) Check the condition of piston rings free movement (Gas sealing)
                                    4) Adjust the Cylinder Oil Feed rate to avoid over lubrication to avoid formation of carbon deposits
                                    """)
                    elif defect[d][2]==k and k=="4":
                        defect[d].append("""1) Replace Piston Rings
                                    2) Check for Carbon deposits in the ring groove
                                    3) Check vertical ring clearance
                                    4) Check for Partial sticking
                                    5) Check for Poor sealing between the ring and the ring groove floor.
                                    6) Check for Clover leafing 
                                    7) Check for Ring end chamfers.
                                    8) Check for too large ring edge radii.
                                    9) Check for Continual striking against wear ridges, or other irregularities in the cylinder wall.
                                    """)
                    elif defect[d][2]==k and k=="3":
                        defect[d].append("""1) Adjust the Cylinder Oil Feed rate
                                    2) Carry out Drain oil analysis (On board or send ashore)
                                    3) Carry out or land samples for Fuel oil analysis
                                    4) Check for Hard abrasive particles
                                    """)             
                    elif defect[d][2]==k and k=="1":
                        defect[d].append("""1) Check Fuel injectors for leakage
                                    2) Check for carbon deposits
                                    """)
                    elif defect[d][2]==k and k=="2":
                        defect[d].append("""1) Adjust the Cylinder Oil Feed rate
                                    2) Carry out Drain oil analysis (On board or send ashore)
                                    3) Adjust feed rate to obtain optimum residual BN
                                    """)
            # print("defect_after_loop:",defect)
            #Overall defect summary
            tensor = {}
            if df["class"].nunique() > 1:
                tensor["multi_defect"]= df.confidence.max()
            else:
                tensor[df.loc[np.argmax(df["confidence"])]['class']] = df.loc[np.argmax(df["confidence"])]['confidence']
            defect.insert(0, tensor)

            #Final object
            # str1,str2 = '',''
            # for item in tensor:
            #     str2 += '('+item + ')-(' + tensor[item] + ')'
            # for l1 in range(1,len(defect)):
            #     str1 = "("+ defect[l1][1]+') - ('+defect[l1][2]+') -'+defect[l1][3]+'-'+defect[l1][4]
            #     str2 += '||'+str1

            # obj={}
            # print("defect:",defect)
            str0,str1 = '',''
            for l1 in range(1,len(defect)):
                str0 = "("+ defect[l1][1]+') - '+defect[l1][3]+' - '+defect[l1][4] + " || " ##Note: Add ring number.
                str1 += str0
            #     print("----------------------",l1)    
            # str1

            obj={}
            
            if df["class"].nunique() > 1:
                obj["Fault_id"] = 99
                obj["Rating"]= df.loc[np.argmax(df["confidence"])]['confidence_lvl']
                obj["Recommendation"] = str1
            else:
                obj["Fault_id"] = df.loc[np.argmax(df["confidence"])]['class']
                obj["Rating"] = df.loc[np.argmax(df["confidence"])]['confidence_lvl']
                obj["Recommendation"] = str1
            # print("obj:",obj)
            obj_lst.append(obj)
            cyl_pred[cyl_number]=obj_lst
    
            # defect.insert(0, tensor)



    ########################################################################################################################################################################################################################

            img_list.append(file_object) ########### adding images in docx #########################
            #print(results.pandas().xyxy[0] ) # im1 predictions (pandas)
            #print("y ",res_tensor[0][1])
            #print("c ",int(res_tensor[0][5]))
            #print("tensor len",len(res_tensor))
            
            rings=[]
            for i in range(0,len(res_tensor)):
                
                #print("percent=====",res_tensor[i][1]/h)
                if res_tensor[i][1]/h <=.25 :
                    rings.append({"1":int(res_tensor[i][5])})
                elif res_tensor[i][1]/h <=.45 :
                    rings.append({"2":int(res_tensor[i][5])})
                elif res_tensor[i][1]/h <=.75 :
                    rings.append({"3":int(res_tensor[i][5])})
                elif res_tensor[i][1]/h >.75 :
                    rings.append({"4":int(res_tensor[i][5])})
        
            def_section_brk=set()
            def_section_lub1=set()
            def_section_surf=set()
            def_section_dep=set()
            def_section_lub2=set()
            def_section_brk_ls={}
            
            def_section_lub_ls={}
            def_section_surf_ls={}
            def_section_dep_ls={}

            # Fault_id added
            
            try:
                for ring_no in range(1,5):
                    
                    def_section_lub_ls["Ring"+str(ring_no)]="*"
            except Exception :
                print("Excepetion")

            try:
                for ring_no in range(1,5):
                    
                    
                    def_section_surf_ls["Ring"+str(ring_no)]="*"
            except Exception :
                print("Excepetion")
            try:
                for ring_no in range(1,5):
                    
                    
                    def_section_dep_ls["Ring"+str(ring_no)]="*"
                    
            except Exception :
                print("Excepetion")
            try:
                for ring_no in range(1,5):
                
                    def_section_brk_ls["Ring"+str(ring_no)]="*"
                
            except Exception :
                print("Excepetion")

            for ring in rings:
                # print(ring.values())
            
                if(list(ring.values())[0]==3):# if collapsed
                    def_section_brk.add(list(ring.keys())[0]) # assign ring number
                if(list(ring.values())[0]==12): 
                    def_section_surf.add(list(ring.keys())[0])
                if(list(ring.values())[0]==11 ):
                    def_section_lub1.add(list(ring.keys())[0])
                if(list(ring.values())[0]==17):
                    def_section_dep.add(list(ring.keys())[0])
                if(list(ring.values())[0]==10):
                    def_section_lub2.add(list(ring.keys())[0])
        
            # for brk in def_section_brk :
            #     #print({"Ring"+brk:"C"})
            #     def_section_brk_ls.update({"Ring"+brk:"C"})
            # for brk in def_section_surf :

            #     def_section_surf_ls.update({"Ring"+brk:"S"})
            # for brk in def_section_dep :
            #     #print({"Ring"+brk:"LC"})
            #     def_section_dep_ls.update({"Ring"+brk:"LC"})

            # for brk in def_section_lub1 :
            #     def_section_lub_ls.update({"Ring"+brk:"OB"})
            # for brk in def_section_lub2 :
            #     #print({"Ring"+brk:"OB"})
            #     if brk not in(list(def_section_lub1)):
            #         def_section_lub_ls.update({"Ring"+brk:"O"})
            # #print(def_section_dep_ls)

            for brk in def_section_brk :
                # #faultdict = {"3"-C: '4',"12"-S: '3',"17"-LC: '5',"10"-O: '2', "11"-OB:'1'}
                #print({"Ring"+brk:faultdict["C"]})
                def_section_brk_ls.update({"Ring"+brk:faultdict["3"]})
            for brk in def_section_surf :
                #print({"Ring"+brk:faultdict["S"]})
                def_section_surf_ls.update({"Ring"+brk:faultdict["12"]})
            for brk in def_section_dep :
                #print({"Ring"+brk:faultdict["LC"]})
                def_section_dep_ls.update({"Ring"+brk:faultdict["17"]})

            for brk in def_section_lub1 :
                #print({"Ring"+brk:faultdict["OB"]})
                def_section_lub_ls.update({"Ring"+brk:faultdict["11"]})
            for brk in def_section_lub2 :
                #print({"Ring"+brk:faultdict["OB"]})
                def_section_lub_ls.update({"Ring"+brk:faultdict["11"]})
                if brk not in(list(def_section_lub1)):
                    def_section_lub_ls.update({"Ring"+brk:faultdict["10"]})
            #print(def_section_dep_ls)
        

            #print("def_section_lub_ls", def_section_lub_ls)
            #js_data=results.pandas().xyxy[0].to_json(orient="records")
            
            defect_df=  {"lubrication":def_section_lub_ls, "surface":def_section_surf_ls,"deposits":def_section_dep_ls,"breakage":def_section_brk_ls
            #, "image":base64img
            }
            print("defect_df line 374:",defect_df)
            # print("data_k",data_k)
            # print("cyl_index",cyl_index)
            defect_df_all_cyl["cylinder"+(data_keys[cyl_index])]=defect_df
            # print("defect_df_all_cyl",defect_df_all_cyl)
            # ##print(defect_df)
            # #print("------------------defect-------------------------------")
        # print("len of tensor:",len_tensor)
    #print(defect_df_all_cyl)

    selection_lubrication=[]
    selection_surface=[]
    selection_deposits=[]
    selection_brekage=[]
        
            
    user_data=[]
            
    cyls=defect_df_all_cyl.keys()
    # print("cyls391",cyls)
        
        
        # for cyl in cyls:
        #     #print(type(cyl))
        #     if cyl.startswith("cylinder") and len_tensor>0 :

        #         selection_lubrication.append(list(defect_df_all_cyl[cyl]['lubrication'].values()))
        

        #         selection_surface.append(list(defect_df_all_cyl[cyl]['surface'].values()))
        

        #         selection_brekage.append(list(defect_df_all_cyl[cyl]['breakage'].values()))
        

        #         selection_deposits.append(list(defect_df_all_cyl[cyl]['deposits'].values()))
        #         print("selection_deposits",selection_deposits,selection_brekage,selection_surface,selection_lubrication)
        #     # elif cyl.startswith("cylinder") and len_tensor==0:
        #     elif cyl
        #         selection_lubrication.append(["*","*","*","*"])
        #         selection_surface.append(["*","*","*","*"])
        #         selection_brekage.append(["*","*","*","*"])
        #         selection_deposits.append(["*","*","*","*"])
        #         print("**************************")
        #         print("selection_deposits",selection_deposits,selection_brekage,selection_surface,selection_lubrication)

    #if len_tensor>0:
    for cyl in cyls:
        #print(type(cyl))
        if cyl.startswith("cylinder"):

            selection_lubrication.append(list(defect_df_all_cyl[cyl]['lubrication'].values()))
    

            selection_surface.append(list(defect_df_all_cyl[cyl]['surface'].values()))
    

            selection_brekage.append(list(defect_df_all_cyl[cyl]['breakage'].values()))
    

            selection_deposits.append(list(defect_df_all_cyl[cyl]['deposits'].values()))
        print("line 435",selection_deposits,selection_brekage,selection_surface,selection_lubrication)
        # elif cyl.startswith("cylinder") and len_tensor==0:
    #else:
    # selection_lubrication.append(["*","*","*","*"])
    # selection_surface.append(["*","*","*","*"])
    # selection_brekage.append(["*","*","*","*"])
    # selection_deposits.append(["*","*","*","*"])
    # print("selection_deposits",selection_deposits,selection_brekage,selection_surface,selection_lubrication)
    json_indep_app={'predictionInfo':[]}
    for lub,sur,dep,bre in zip(selection_lubrication,selection_surface,selection_deposits,selection_brekage):
        for i,j,k,l in zip(lub,sur,dep,bre):
            json_indep_app['predictionInfo'].append({"lubricationCondition":i,"surfaceCondition":j,"depositsCondition":k,"breakageCondition":l})
    print("json_indep_app line 447",json_indep_app)



    pred_per_cyl_lubrication_rev = [[selection_lubrication[j][i] for j in range(len(selection_lubrication))] for i in range(len(selection_lubrication[0]))]
    pred_per_cyl_surface_rev = [[selection_surface[j][i] for j in range(len(selection_surface))] for i in range(len(selection_surface[0]))]
    pred_per_cyl_deposits_rev = [[selection_deposits[j][i] for j in range(len(selection_deposits))] for i in range(len(selection_deposits[0]))]
    pred_per_cyl_breakage_rev = [[selection_brekage[j][i] for j in range(len(selection_brekage))] for i in range(len(selection_brekage[0]))]
    #print(pred_per_cyl_lubrication_rev)
            
    
    
    # for i,j in zip(data_keys,obj_lst): ##Final Object
    #     cyl_obj[i]=j
    # final_cyl_obj={**cyl_obj, **data_id}
    # print("Object:", final_cyl_obj)
    # result_obj = json.loads(json.dumps(final_cyl_obj))  
    # print(result_obj)

    data_pdf={'PDF':"https://scavaiapp.azurewebsites.net/files/report.pdf"} 
    # data_pdf={'PDF':"https://localhost:3000/files/report.pdf"} 
    # data_pdf['PDF']= [t.encode('utf-8') for t in title]
    bytes_io_object=img_list[0]
    # Get the binary data from the BytesIO object
    binary_image_data = bytes_io_object.getvalue()

    # Convert the binary data to base64 encoding
    base64_encoded_image = base64.b64encode(binary_image_data).decode('utf-8')
    json_indep_app['image']=base64_encoded_image
    print("json_indep_app line1150",json_indep_app)
    for i,j in zip(data_k,obj_lst): ##Final Object
        cyl_obj[i]=j
    final_cyl_obj={**cyl_obj,**data_pdf, **data_id}
    # print("data_k:",data_k)
    # print("obj_lst:",obj_lst)
    # print("cyl_pred:",cyl_pred)
    # print("final_cyl_obj:",final_cyl_obj)
    # print("Object:", final_cyl_obj)
    result_obj = json.loads(json.dumps(json_indep_app))  
    # print(result_obj)
    return result_obj










def excel(data):
    print(" @@@@ DOWNLOAD EXCEL")
    
    
    selection_lubrication=[]
    selection_surface=[]
    selection_deposits=[]
    selection_brekage=[]
    cylinder_cells=[]
    user_data=[]
    
    pred_json=data['predictionInfo']
    pred_json = {key: value for key, value in pred_json.items() if value is not None}
    user_data=(data['info'])
    cyl_max=0
    for cyl in pred_json:
        
        if cyl.startswith("cylinder"):
            if len(cyl)>9:
                cyl_num=int(cyl[len(cyl)-2:])
                print(cyl_num)
            else:
                    cyl_num=int(cyl[-1])
            if cyl_num>cyl_max:
                cyl_max=cyl_num
    excel_cyl_columns=[]
    
    cyl_order_list= list(range(1,cyl_max+1))
    print("MAXXXX==",cyl_order_list ,pred_json.keys())
        ##################73 ---- ASCII I ===== Making the columns for dynamic cylinders like I,J,K,L,M,N etc
    for ii in cyl_order_list:
        excel_cyl_columns.append(chr(72+(ii) ))
    print(" Excel Cylinder Cloumns", excel_cyl_columns)

    no_fault=['N/A', 'N/A', 'N/A', 'N/A']
    for ord in cyl_order_list:
        cyl="cylinder"+str(ord)
        if cyl in(pred_json.keys()):
            
            print(cyl)
            print(list(pred_json[cyl]['lubrication'].values()))
            selection_lubrication.append(list(pred_json[cyl]['lubrication'].values()))
    

            selection_surface.append(list(pred_json[cyl]['surface'].values()))
    

            selection_brekage.append(list(pred_json[cyl]['breakage'].values()))
    

            selection_deposits.append(list(pred_json[cyl]['deposit'].values()))
        else:
            selection_lubrication.append(list(no_fault))
    

            selection_surface.append(list(no_fault))
    

            selection_brekage.append(list(no_fault))
    

            selection_deposits.append(list(no_fault))
   
    pred_per_cyl_lubrication_rev = [[selection_lubrication[j][i] for j in range(len(selection_lubrication))] for i in range(len(selection_lubrication[0]))]
    pred_per_cyl_surface_rev = [[selection_surface[j][i] for j in range(len(selection_surface))] for i in range(len(selection_surface[0]))]
    pred_per_cyl_deposits_rev = [[selection_deposits[j][i] for j in range(len(selection_deposits))] for i in range(len(selection_deposits[0]))]
    pred_per_cyl_breakage_rev = [[selection_brekage[j][i] for j in range(len(selection_brekage))] for i in range(len(selection_brekage[0]))]
    print(pred_per_cyl_lubrication_rev )
    print(pred_per_cyl_surface_rev)
    print(pred_per_cyl_deposits_rev)
    print(pred_per_cyl_breakage_rev)
     ########################################  EXCEL #########################################
    #local_path = os.path.join(tempfile.gettempdir(), 'report.xlsx')
    wb = openpyxl.load_workbook('new_sample_demo.xlsx')
    #print(" local_path = os.path.join(tempfile.gettempdir(), 'prediction.xlsx')")
    print("USER+++++++++++++++++++++",user_data)
    ws = wb.worksheets[0]
    ws['E7']=user_data['vessel_name'] 

    #'company_name': 'Treeswise', 'vessel_name': 'vessel Name', 'imo_number': '1232323', 
    # 'manufacturer': 'manufacturer', 'type_of_engine': 'type of engine', 'vessel_type': 'vessel type', 
    # 'inspection_date': '2023-03-28', 'total_running_hours': '3427', 'running_hrs_since_last': '23',
    #  'cyl_oil_Type': '759387', 'cyl_oil_consump_Ltr_24hr': '878', 
    # 'normal_service_load_in_percent_MCR': '8738', 'cylinder_numbers': '10' 
    ws['E8']=user_data['imo_number']
    ws['E9']=user_data['manufacturer'] 
    ws['E10']=user_data['type_of_engine'] 
    ws['E11']=user_data['vessel_type']  
    ws['E12']=user_data['cylinder_numbers']
    
    ws['M7']=pred_json['date'] 
    ws['M8']=pred_json['total_running_hours'] 
    ws['M9']=user_data['running_hrs_since_last'] 
    ws['M10']=user_data['cyl_oil_Type'] 
    ws['M11']=user_data['cyl_oil_consump_Ltr_24hr'] 
    ws['M12']=user_data['normal_service_load_in_percent_MCR']  
    cyl_no_col=1 # WS[I15] =1 , WS[J15] =2 , etc
    for j in (excel_cyl_columns):
         ws[j+str(15) ]= cyl_no_col
         cyl_no_col=cyl_no_col+1
    for i in range(4):
        excelcellind=0
        for j in (excel_cyl_columns):
            #print (i,j,excelcellind ,"     ",j+str(20+i),"   ",  pred_per_cyl_deposits[i][excelcellind])
            ws[j+str(17+i) ] = pred_per_cyl_deposits_rev[i][excelcellind]
            excelcellind+=1
    for i in range(4):
        excelcellind=0
        for j in (excel_cyl_columns):
            #print (i,j,excelcellind ,"     ",j+str(20+i),"   ",  pred_per_cyl_surface[i][excelcellind])
            ws[j+str(27+i) ] = pred_per_cyl_surface_rev[i][excelcellind]
            excelcellind+=1
    for i in range(4):
        excelcellind=0
        for j in (excel_cyl_columns):
            # print (i,j,excelcellind ,"     ",j+str(24+i),"   ",  pred_per_cyl_lubrication_rev[i][excelcellind])
            ws[j+str(32+i) ] = pred_per_cyl_lubrication_rev[i][excelcellind]
            excelcellind+=1
    for i in range(4):
        excelcellind=0
        for j in (excel_cyl_columns):
            # print (i,j,excelcellind ,"     ",j+str(24+i),"   ",  pred_per_cyl_lubrication_rev[i][excelcellind])
            ws[j+str(22+i) ] = pred_per_cyl_breakage_rev[i][excelcellind]
            excelcellind+=1
    unit_no=0
    for cyl in pred_json:
            
            if cyl.startswith("cylinder"):
               ws["A"+str(38+unit_no)]=cyl
               ws["B"+str(38+unit_no)]=pred_json[cyl]['remark']
               unit_no=unit_no+1
    
    ws = wb.worksheets[1]
    imgexcelwidth=0
    for cyl in (pred_json):
       if cyl.startswith("cylinder"):
            index=1
            print("--------------Before Image saving ---")
            try:
            
                encoded = pred_json[cyl]['image']#. split(",", 1) #.decode("utf-8")
                data_img=base64.b64decode(encoded)
                
                img_path = "image"+cyl+".png"
                with open(img_path, "wb") as f:
                    f.write(data_img)
                
                img_stream = Image.open(io.BytesIO(data_img))
                #data_PIL = Image.fromarray(img)
                
                
                img_h=img_stream.height
            
                img = openpyxl.drawing.image.Image(img_path)
                
                img_height_line=int(7*img_h/100)
            
                
                ind=(imgexcelwidth+2) 
                ws.add_image(img ,'D'+str(ind))
                imgexcelwidth=imgexcelwidth+img_height_line
                
            except Exception as e:
            #     print('exception',e.__class__ )
                print(str(e))
                continue
       
    wb.save('./files/prediction.xlsx')
    wb.close()
    wb = None
       
    
    #data_pdf={'PDF':"https://scavaiapp.azurewebsites.net/files/report.pdf"} 

   
    file='./files/prediction.xlsx'

    
    return FileResponse(file, media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")

    
def pdf(data):
    
    
    selection_lubrication=[]
    selection_surface=[]
    selection_deposits=[]
    selection_brekage=[]
    cyl_remarks=[]
    user_data=[]
    img_list=[]
    
    
    cyl_pred={}
    pred_json=data['predictionInfo']
    pred_json = {key: value for key, value in pred_json.items() if value is not None}
    user_data=(data['info'])
    for i,j in pred_json.items():
        if i.startswith("cylinder"):
            cyl_pred[i]=j
    myKeys = list(cyl_pred.keys())
    print("ddddddddddd",myKeys)
    cyl_num_temp=[re.split(r'(\d+)', s) for s in myKeys]
    cyl_num=[int(i[1]) for i in cyl_num_temp]
    Z = [x for _,x in sorted(zip(cyl_num,myKeys))]
    cyl_pred_sorted={}
    cyl_num.sort()
    for i in Z:
        cyl_pred_sorted[i]=cyl_pred[i]



    for cyl in cyl_pred_sorted:
        print("ssssssss",cyl)
        #header, encoded = cyl_pred_sorted[cyl]['image'].split(",", 1)
        encoded = cyl_pred_sorted[cyl]['image']
        data = base64.b64decode(encoded)

        with open("image.png", "wb") as f:
            f.write(data)   


    
    
        img_list.append(io.BytesIO(data))
        
        
        print(list(cyl_pred_sorted[cyl]['lubrication'].values()))
        selection_lubrication.append(list(cyl_pred_sorted[cyl]['lubrication'].values()))
        selection_surface.append(list(cyl_pred_sorted[cyl]['surface'].values()))
        selection_brekage.append(list(cyl_pred_sorted[cyl]['breakage'].values()))
        selection_deposits.append(list(cyl_pred_sorted[cyl]['deposit'].values()))
        cyl_remarks.append(cyl_pred_sorted[cyl]['remark'])
    pred_per_cyl_lubrication_rev = [[selection_lubrication[j][i] for j in range(len(selection_lubrication))] for i in range(len(selection_lubrication[0]))]
    pred_per_cyl_surface_rev = [[selection_surface[j][i] for j in range(len(selection_surface))] for i in range(len(selection_surface[0]))]
    pred_per_cyl_deposits_rev = [[selection_deposits[j][i] for j in range(len(selection_deposits))] for i in range(len(selection_deposits[0]))]
    pred_per_cyl_breakage_rev = [[selection_brekage[j][i] for j in range(len(selection_brekage))] for i in range(len(selection_brekage[0]))]
    print(pred_per_cyl_lubrication_rev )
    print(pred_per_cyl_surface_rev)
    print(pred_per_cyl_deposits_rev)
    print(pred_per_cyl_breakage_rev)
    print(cyl_num)
    print("=====================================",pred_json.keys())

    doc = Document('sample_report.docx')
    doc.tables #a list of all tables in document
   


    doc.tables[0].cell(0, 1).text = user_data["vessel_name"] # vessel name
    doc.tables[0].cell(0, 1).paragraphs[0].runs[0].font.size = Pt(7)
    doc.tables[0].cell(0, 1).paragraphs[0].runs[0].font.name = 'Arial'
    doc.tables[0].cell(0, 1).paragraphs[0].runs[0].font.bold = True

    # doc.tables[0].cell(0, 3).text = user_data["hullNubmer"] # hull
    # doc.tables[0].cell(0, 3).paragraphs[0].runs[0].font.size = Pt(7)
    # doc.tables[0].cell(0, 3).paragraphs[0].runs[0].font.name = 'Arial'
    # doc.tables[0].cell(0, 3).paragraphs[0].runs[0].font.bold = True

    doc.tables[0].cell(0, 5).text=user_data["vessel_type"] # vessel type
    doc.tables[0].cell(0,5).paragraphs[0].runs[0].font.size = Pt(7)
    doc.tables[0].cell(0,5).paragraphs[0].runs[0].font.name = 'Arial'
    doc.tables[0].cell(0,5).paragraphs[0].runs[0].font.bold = True

    doc.tables[2].cell(0, 1).text = user_data["manufacturer"] # make
    doc.tables[2].cell(0, 1).paragraphs[0].runs[0].font.size = Pt(7)
    doc.tables[2].cell(0, 1).paragraphs[0].runs[0].font.name = 'Arial'
    doc.tables[2].cell(0, 1).paragraphs[0].runs[0].font.bold = True
    
    doc.tables[4].cell(1, 1).text = str(pred_json["total_running_hours"]) #  Total running hours
    doc.tables[4].cell(1, 1).paragraphs[0].runs[0].font.size = Pt(7)
    doc.tables[4].cell(1, 1).paragraphs[0].runs[0].font.name = 'Arial'


    doc.tables[4].cell(2, 1).text = user_data["cyl_oil_Type"] # Cyl oil type
    doc.tables[4].cell(2, 1).paragraphs[0].runs[0].font.size = Pt(7)
    doc.tables[4].cell(2, 1).paragraphs[0].runs[0].font.name = 'Arial'
   

    doc.tables[4].cell(3, -1).text = user_data["cyl_oil_consump_Ltr_24hr"] # Cyl oil rate
    doc.tables[4].cell(3, -1).paragraphs[0].runs[0].font.size = Pt(7)
    doc.tables[4].cell(3, -1).paragraphs[0].runs[0].font.name = 'Arial'
   

    
    doc.tables[4].cell(3, 1).text = str(user_data["normal_service_load_in_percent_MCR"]) # Normal service MCR
    doc.tables[4].cell(3, 1).paragraphs[0].runs[0].font.size = Pt(7)
    doc.tables[4].cell(3, 1).paragraphs[0].runs[0].font.name = 'Arial'

#     lubrication_regulation
    
    
    
    
#     totalRunningHours
#     vesselName
#     vesselType
    
   
    list_ind=[2,3,-2,-1]
    doc.tables[5].style ='TableGrid' # Deposits section --tables[5]
    for i,k in zip(range(len(selection_deposits)),cyl_num):
        row_cells = doc.tables[5].add_row().cells
        ind=0
        row_cells[0].text = str(k)
        row_cells[0].paragraphs[0].runs[0].font.bold = True
        row_cells[0].paragraphs[0].runs[0].font.size = Pt(7)
        row_cells[0].paragraphs[0].runs[0].font.name = 'Arial'
        for j in list_ind:
            row_cells[j].text = 'l' if selection_deposits[i][ind]=='*' else selection_deposits[i][ind]
           # row_cells[0].paragraphs[0].runs[0].font.bold = True
            row_cells[j].paragraphs[0].runs[0].font.size = Pt(7)
            row_cells[j].paragraphs[0].runs[0].font.name = 'Arial'
            ind=ind+1
        row_cells[1].text = 'l' 
        # row_cells[0].paragraphs[0].runs[0].font.bold = True
        row_cells[1].paragraphs[0].runs[0].font.size = Pt(7)
        row_cells[1].paragraphs[0].runs[0].font.name = 'Arial'
    #Mean row
    # row_cells = doc.tables[5].add_row().cells
    # row_cells[0].text = "Mean"
    # row_cells[0].paragraphs[0].runs[0].font.bold = True
    # row_cells[0].paragraphs[0].runs[0].font.size = Pt(7)
    # row_cells[0].paragraphs[0].runs[0].font.name = 'Arial'
    # Breakage section --tables[6]
    list_ind=[1,2,3,-2]
    doc.tables[6].style ='TableGrid' 
    for i,k in zip(range(len(selection_brekage)),cyl_num):
        row_cells = doc.tables[6].add_row().cells
        ind=0
        row_cells[0].text = str(k)
        row_cells[0].paragraphs[0].runs[0].font.bold = True
        row_cells[0].paragraphs[0].runs[0].font.size = Pt(7)
        row_cells[0].paragraphs[0].runs[0].font.name = 'Arial'
        for j in list_ind:
            row_cells[j].text ='l' if selection_brekage[i][ind]=='*' else selection_brekage[i][ind]
           # row_cells[0].paragraphs[0].runs[0].font.bold = True
            row_cells[j].paragraphs[0].runs[0].font.size = Pt(7)
            row_cells[j].paragraphs[0].runs[0].font.name = 'Arial'
            ind=ind+1
        
        row_cells[-1].text = 'l' 
        # row_cells[0].paragraphs[0].runs[0].font.bold = True
        row_cells[-1].paragraphs[0].runs[0].font.size = Pt(7)
        row_cells[-1].paragraphs[0].runs[0].font.name = 'Arial'

    # row_cells = doc.tables[6].add_row().cells
    # row_cells[0].text = "Mean"
    # row_cells[0].paragraphs[0].runs[0].font.bold = True
    # row_cells[0].paragraphs[0].runs[0].font.size = Pt(7)
    # row_cells[0].paragraphs[0].runs[0].font.name = 'Arial'

# Surface section --tables[7]
    list_ind=[1,2,3,4]
    list_ind_extra=[5,6,7,-2,-1]
    doc.tables[7].style ='TableGrid' 
    for i,k in zip(range(len(selection_surface)),cyl_num):
        row_cells = doc.tables[7].add_row().cells
        ind=0
        row_cells[0].text = str(k)
        row_cells[0].paragraphs[0].runs[0].font.bold = True
        row_cells[0].paragraphs[0].runs[0].font.size = Pt(7)
        row_cells[0].paragraphs[0].runs[0].font.name = 'Arial'
        for j in list_ind:
            row_cells[j].text = 'Cl' if selection_surface[i][ind]=='*' else selection_surface[i][ind]
           # row_cells[0].paragraphs[0].runs[0].font.bold = True
            row_cells[j].paragraphs[0].runs[0].font.size = Pt(7)
            row_cells[j].paragraphs[0].runs[0].font.name = 'Arial'
            ind=ind+1
        for ext in list_ind_extra:
            row_cells[ext].text = 'Cl' 
           # row_cells[0].paragraphs[0].runs[0].font.bold = True
            row_cells[ext].paragraphs[0].runs[0].font.size = Pt(7)
            row_cells[ext].paragraphs[0].runs[0].font.name = 'Arial'
        
    # row_cells = doc.tables[7].add_row().cells #Mean row
    # row_cells[0].text = "Mean"
    # row_cells[0].paragraphs[0].runs[0].font.bold = True
    # row_cells[0].paragraphs[0].runs[0].font.size = Pt(7)
    # row_cells[0].paragraphs[0].runs[0].font.name = 'Arial'

# Lubrication section --tables[8]
    list_ind=[1,2,3,4]
    list_ind_extra=[5,6,-2,-1]
    doc.tables[8].style ='TableGrid' 
    for i,k in zip(range(len(selection_lubrication)),cyl_num):
        row_cells = doc.tables[8].add_row().cells
        ind=0
        row_cells[0].text = str(k)
        row_cells[0].paragraphs[0].runs[0].font.bold = True
        row_cells[0].paragraphs[0].runs[0].font.size = Pt(7)
        row_cells[0].paragraphs[0].runs[0].font.name = 'Arial'
        for j in list_ind:
            row_cells[j].text = 'N' if selection_lubrication[i][ind]=='*' else selection_lubrication[i][ind]
           # row_cells[0].paragraphs[0].runs[0].font.bold = True
            row_cells[j].paragraphs[0].runs[0].font.size = Pt(7)
            row_cells[j].paragraphs[0].runs[0].font.name = 'Arial'
            ind=ind+1
        
        for ext in list_ind_extra: # Ring5 , Piston Skirt,Rod,Liner
            row_cells[ext].text = 'N' 
           # row_cells[0].paragraphs[0].runs[0].font.bold = True
            row_cells[ext].paragraphs[0].runs[0].font.size = Pt(7)
            row_cells[ext].paragraphs[0].runs[0].font.name = 'Arial'

    # row_cells = doc.tables[8].add_row().cells  #Mean row
    # row_cells[0].text = "Mean"
    # row_cells[0].paragraphs[0].runs[0].font.bold = True
    # row_cells[0].paragraphs[0].runs[0].font.size = Pt(7)
    # row_cells[0].paragraphs[0].runs[0].font.name = 'Arial'

# Remarks section --tables[9]
    list_ind=[1]
    doc.tables[9].style ='TableGrid' 
    for i,k in zip(range(len(cyl_num)),cyl_num):
        row_cells = doc.tables[9].add_row().cells
        ind=0
        row_cells[0].text = str(k)
        row_cells[0].paragraphs[0].runs[0].font.bold = True
        row_cells[0].paragraphs[0].runs[0].font.size = Pt(7)
        row_cells[0].paragraphs[0].runs[0].font.name = 'Arial'
        for j in list_ind:
            row_cells[j].text = cyl_remarks[i]
           # row_cells[0].paragraphs[0].runs[0].font.bold = True
            row_cells[j].paragraphs[0].runs[0].font.size = Pt(7)
            row_cells[j].paragraphs[0].runs[0].font.name = 'Arial'
            ind=ind+1
        
    # row_cells = doc.tables[9].add_row().cells #Mean row
    # row_cells[0].text = "Mean"
    # row_cells[0].paragraphs[0].runs[0].font.bold = True
    # row_cells[0].paragraphs[0].runs[0].font.size = Pt(7)
    # row_cells[0].paragraphs[0].runs[0].font.name = 'Arial'


    ########### adding images #########################
    
 
    table = doc.add_table(rows=math.ceil(len(img_list)/2), cols=2)
    table.style ='TableGrid' 

    im_count=0
    im_id=[]
    for i in cyl_num:
        im_id.append(int(i))  
    print(len(img_list))

    for row,im_num in zip(table.rows,im_id):
        for cell in row.cells:
            im_count+=1
            
            if im_count<= len(img_list): # odd number of cylinders
                cell.text = "Cylinder "+str(im_id[im_count-1])
                p=cell.add_paragraph()
                r = p.add_run()
                print("-------",img_list[im_count-1])

                r.add_picture(img_list[im_count-1],width=Inches(3))




    doc.save("./report.docx")

    # !pip install docx2pdf
    # convert docx to pdf
    # using docx2pdf module
    
    # pythoncom.CoInitialize()

    # as the python file
#     convert("report.docx")
    
    # output = convert(source="report.docx", output_dir="./", soft=1)
    #command = "abiword -t report_output.pdf report.docx"
    #p = Popen(command)
    #p.communicate()
    args = ["abiword", "--to", "report_output.pdf", "report.docx"  ]

    call(args )

    #check_call([ 'abiword', '-t ','report_output.pdf','report.docx'], stdout=open(os.devnull,'wb'), stderr=STDOUT)
    file='report_output.pdf'

    
    return FileResponse(file, media_type="application/pdf")

    
   
 
